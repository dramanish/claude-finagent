#!/usr/bin/env python3
"""
QSR Digest — research + classify + Word doc
Outputs latest-classified.json for the Claude Code cron to pick up.
"""

import os
import re
import json
from datetime import date
from pathlib import Path
import anthropic

client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
KB = Path(__file__).parent / "knowledge-base"
OUT = Path(__file__).parent / "out"
OUT.mkdir(exist_ok=True)

MODEL = "claude-opus-4-5-20251101"


def extract_json(text):
    # Try JSON code block first
    block = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if block:
        return json.loads(block.group(1))
    # Fall back to outermost braces
    start, end = text.find("{"), text.rfind("}") + 1
    return json.loads(text[start:end])


def read_kb(*filenames):
    parts = []
    for fn in filenames:
        path = KB / fn
        if path.exists():
            parts.append(f"## {fn}\n{path.read_text()}")
    return "\n\n".join(parts)


def update_search_history(items):
    history_path = KB / "Search History.md"
    today = str(date.today())
    new_lines = "\n".join(
        f"- {today} | {item['headline']} | {item['source_url']}"
        for item in items
    )
    existing = history_path.read_text()
    if f"## {today}" in existing:
        updated = existing + "\n" + new_lines
    else:
        updated = existing + f"\n\n## {today}\n{new_lines}"
    history_path.write_text(updated)
    print(f"  Search history updated ({len(items)} items logged)")


def run_sector_reader():
    print("\n[1/3] Sector reader — researching public QSR sources...")
    context = read_kb(
        "Research Priorities.md", "Daniels Donuts.md",
        "Krispy Kreme AU.md", "Donut King.md",
        "Collins Foods.md", "Brooklyn Donuts.md"
    )
    history = read_kb("Search History.md")
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        tools=[{"type": "web_search_20250305", "name": "web_search"}],
        system=f"""You are the research subagent for LK Group / Daniels Donuts.
Your job is targeted intelligence gathering — not generic news monitoring.

## Search protocol — run at least 8 targeted searches
Work through each standing priority systematically:

1. "Krispy Kreme Australia" + current month/year — channel deals, store openings, partnerships
2. "Donut King" OR "Retail Food Group" ASX — results, digital updates, strategy
3. "Collins Foods" CKF ASX — latest results, same-store sales, trading updates
4. "Brooklyn Donuts" Australia — store openings, franchise news, VIC expansion
5. Australia QSR fast food labour wage award 2025 OR 2026 — Fair Work, award rates
6. "Uber Eats" OR "DoorDash" Australia fees commission 2025 OR 2026
7. Australia donut specialty new brand OR entrant 2025 OR 2026
8. "LK Group" OR "Queens Lane Capital" OR "Daniels Donuts" — portfolio news

Run extra searches if any priority returns no fresh results.

## Source standards
- Public only: ASX announcements, news sites, LinkedIn, public broker research
- Prefer sources from the last 30 days
- Items older than 30 days only if major (M&A, regulatory change, significant result)
- Every item MUST have a real source URL — no URL, no item
- Cross-reference significant claims across 2 sources where possible

## Deduplication
Check every item against Search History before including it.
If the headline or source URL already appears — skip it entirely.

## Output format
Return a JSON object with this exact structure:
{{
  "items": [
    {{
      "headline": "...",
      "summary": "2-3 sentences",
      "source_url": "...",
      "source_name": "...",
      "date": "YYYY-MM-DD",
      "companies_mentioned": ["..."]
    }}
  ]
}}

Return only items that are genuinely new and newsworthy. Quality over quantity — 2 sharp items beats 8 vague ones. JSON only — no other text.

--- KNOWLEDGE BASE ---
{context}

--- SEARCH HISTORY (skip anything already here) ---
{history}""",
        messages=[{"role": "user", "content": "Run the weekly QSR research scan now."}]
    )
    text = "".join(b.text for b in response.content if hasattr(b, "text"))
    return extract_json(text)


def run_classifier(research):
    print("\n[2/3] Classifier — labelling items GM / Board / Both...")
    context = read_kb("GM vs Board Guide.md", "Daniels Donuts.md", "Research Priorities.md")
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=f"""You are the audience classifier for the LK Group QSR daily digest.

## Source credibility check — apply FIRST before anything else
LK Group is a serious business. Only include items from sources they would act on.

**Tier 1 — Accept:** ASX announcements (asx.com.au), Australian Financial Review,
The Australian, Sydney Morning Herald, Reuters, Bloomberg, Wall Street Journal,
Financial Times, Inside Retail Australia, QSR Media, Restaurant Dive,
Fair Work Commission / government bodies (fairwork.gov.au, abs.gov.au),
official company investor relations pages, official company LinkedIn pages,
public broker research (Bell Potter, Morgans, Macquarie, UBS, Citi, etc.)

**Tier 2 — Flag:** Trade publications, industry newsletters, regional news outlets,
reputable food industry blogs with named authors and editorial standards.
Include but add a note "(TRADE SOURCE)" to the source_name field.

**Tier 3 — Reject:** Anonymous blogs, unknown websites, social media posts
from non-official accounts, user-generated content, sites with no identifiable
editorial standards, press release aggregators with no original reporting.
Do not include these items at all — remove them from the output entirely.

Rule of thumb: would the AFR cite this source? If no — reject it.

## For each item that passes the source check, assign:
- audience: "GM", "Board", or "Both" — use the GM vs Board Guide strictly
- why_it_matters: one sentence, Daniels Donuts specific. Name the implication not
  the observation. Bad: "This shows the QSR sector is competitive." Good: "Krispy
  Kreme's bp deal adds a 6th impulse-purchase channel — directly erodes Daniels'
  grab-and-go advantage in petrol/convenience."
- priority: "HIGH" if it directly matches one of the 8 standing research priorities,
  "NORMAL" otherwise

## Cut rule — apply before classifying
If an item would not change how the GM runs the business or how the Board
thinks about risk or strategy — cut it entirely. Do not include it.
Less is more. 2 sharp items is a better digest than 8 vague ones.

## Rigour checks before classifying each item:
1. Does the source pass Tier 1 or 2? If Tier 3 — exclude entirely
2. Would this item change a decision or prompt an action? If no — cut it
3. Is the "why it matters" specific to Daniels or generic sector commentary?
   If you can't write a Daniels-specific implication — cut the item
4. Does audience match the GM vs Board Guide exactly? When in doubt — Both

Return a JSON object:
{{
  "classified_items": [
    {{
      "headline": "...",
      "summary": "...",
      "source_url": "...",
      "source_name": "...",
      "date": "...",
      "audience": "GM|Board|Both",
      "why_it_matters": "...",
      "priority": "HIGH|NORMAL"
    }}
  ]
}}

JSON only — no other text.

--- KNOWLEDGE BASE ---
{context}""",
        messages=[{
            "role": "user",
            "content": f"Classify these research items:\n{json.dumps(research, indent=2)}"
        }]
    )
    text = "".join(b.text for b in response.content if hasattr(b, "text"))
    return extract_json(text)


def run_note_writer(classified):
    print("\n[3/3] Note writer — producing Word document...")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("LK Group — QSR Intelligence Digest", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    items_list = classified["classified_items"]
    high_count = sum(1 for i in items_list if i["priority"] == "HIGH")

    meta = doc.add_paragraph()
    meta.add_run(f"Week of {date.today().strftime('%d %B %Y')}").bold = True
    doc.add_paragraph(f"Prepared for: General Manager / Board  |  Daniels Donuts")
    doc.add_paragraph(f"{len(items_list)} item{'s' if len(items_list) != 1 else ''} — {high_count} HIGH priority")
    doc.add_paragraph("Sources: Public — ASX filings, news, public broker research")
    doc.add_paragraph("")

    order = {"HIGH": 0, "NORMAL": 1}
    audience_order = {"Both": 0, "Board": 1, "GM": 2}
    items = sorted(
        classified["classified_items"],
        key=lambda x: (order[x["priority"]], audience_order[x["audience"]])
    )

    for item in items:
        label = f"[{item['audience']}]"
        if item["priority"] == "HIGH":
            label += "  ★ HIGH PRIORITY"
        doc.add_heading(item["headline"], level=2)
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        doc.add_paragraph(item["summary"])
        why = doc.add_paragraph()
        why.add_run("Why it matters: ").bold = True
        why.add_run(item["why_it_matters"])
        src = doc.add_paragraph()
        src.add_run("Source: ").bold = True
        src.add_run(f"{item['source_name']} — {item['source_url']}")
        doc.add_paragraph("")

    out_path = OUT / f"qsr-digest-{date.today()}.docx"
    doc.save(out_path)
    return out_path, items


def main():
    print("=" * 60)
    print("  QSR Digest Agent — LK Group / Daniels Donuts")
    print("=" * 60)

    research = run_sector_reader()
    raw_count = len(research.get("items", []))
    print(f"  Found {raw_count} raw items")

    if raw_count == 0:
        print("\n  No items found this run — nothing to digest. Exiting.")
        print("=" * 60)
        return

    classified = run_classifier(research)
    items_after = classified.get("classified_items", [])
    print(f"  {len(items_after)} items passed source + relevance checks (from {raw_count} raw)")

    if not items_after:
        print("\n  All items were cut by the classifier (low quality or irrelevant). No digest sent.")
        print("=" * 60)
        return

    docx_path, items = run_note_writer(classified)
    update_search_history(items)

    summary = {
        "date": str(date.today()),
        "docx_path": str(docx_path),
        "item_count": len(items),
        "items": items
    }
    summary_path = Path(__file__).parent / "out" / "latest-classified.json"
    summary_path.write_text(json.dumps(summary, indent=2))

    print(f"\n  Digest saved: {docx_path}  ({len(items)} item{'s' if len(items) != 1 else ''})")
    print(f"  Summary JSON: {summary_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
